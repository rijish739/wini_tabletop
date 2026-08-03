from __future__ import annotations

"""Build the child-safe interaction and safeguarding specification.

This source is intentionally kept with the artifact so the team can revise the
document without retyping formatted content.  The document is created with the
bundled python-docx runtime and the skill's exact table-geometry helper.
"""

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).resolve().parent / "AI_Tutor_Child_Safe_Interaction_Specification.docx"
SKILL_SCRIPTS = Path(
    r"C:\Users\LENOVO\.codex\plugins\cache\openai-primary-runtime\documents\26.723.12215\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "595959"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
CAUTION = "FFF4CE"
RISK = "FCE8E6"
GREEN = "E2F0D9"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_text(cell, text, *, bold=False, color="000000", size=10.2):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Calibri")
    font.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(font)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.append(size)
    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def format_paragraph(p, *, before=0, after=6, line=1.10, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    style_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in style_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("AI TUTOR - CHILD-SAFE INTERACTION SPECIFICATION")
    set_run_font(r, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    r = footer.add_run("Internal product specification | v1.0 | 28 July 2026")
    set_run_font(r, size=8.5, color=MUTED)


def add_title(doc):
    p = doc.add_paragraph()
    format_paragraph(p, before=34, after=4, line=1.0)
    r = p.add_run("AI Tutor: Child-Safe Interaction and Safeguarding Specification")
    set_run_font(r, size=24, color=INK, bold=True)
    p = doc.add_paragraph()
    format_paragraph(p, after=18, line=1.1)
    r = p.add_run("How the tutor should handle confusion, incomplete speech, off-topic requests, personal data, and emotional or safety disclosures")
    set_run_font(r, size=13, color=MUTED)

    p = doc.add_paragraph()
    shade_paragraph(p, CAUTION)
    format_paragraph(p, before=4, after=4, line=1.08)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    r = p.add_run("Decision status: safety-critical internal guidance. This document sets product requirements and example language; it is not legal advice, medical advice, or a substitute for trained safeguarding staff.")
    set_run_font(r, size=10.5, color="7A5A00", bold=True)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, *, italic=False):
    p = doc.add_paragraph()
    format_paragraph(p)
    r = p.add_run(text)
    set_run_font(r, size=11, italic=italic)
    return p


def add_labelled(doc, label, text, *, color=INK):
    p = doc.add_paragraph()
    format_paragraph(p, after=5)
    r = p.add_run(label)
    set_run_font(r, size=11, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, size=11)
    return p


def add_callout(doc, title, text, *, fill=CALLOUT, color=INK):
    p = doc.add_paragraph()
    shade_paragraph(p, fill)
    format_paragraph(p, before=4, after=4, line=1.08)
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    r = p.add_run(title + " ")
    set_run_font(r, size=10.6, color=color, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.6, color="000000")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, *, header_fill=LIGHT_GRAY, font_size=9.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        set_cell_text(cell, h, bold=True, color=INK, size=font_size)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=font_size)
    apply_table_geometry(table, widths, table_width_dxa=9360, indent_dxa=120)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_quote(doc, text):
    p = doc.add_paragraph()
    shade_paragraph(p, "F9FAFB")
    format_paragraph(p, before=4, after=4, line=1.08)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    r = p.add_run('"' + text + '"')
    set_run_font(r, size=10.6, color=INK, italic=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def build_document():
    doc = Document()
    configure_document(doc)
    add_title(doc)

    add_heading(doc, "1. Executive decision", 1)
    add_body(
        doc,
        "The tutor must behave as a learning system first, a conversational system second, and a safeguarding system whenever a child may be at risk. It must not guess a topic, label a child by IQ, treat speech-recognition text as certain, or use a generic reassuring reply for every safety disclosure. The safe default is: preserve dignity, make the smallest useful next move, ask for clarification when meaning is uncertain, and hand off to a designed human-support pathway when education is no longer the primary need.",
    )
    add_callout(
        doc,
        "Launch red line:",
        "Do not launch crisis or trauma handling as a feature unless there is a documented, staffed, locale-aware escalation pathway; clear product copy that does not promise secrecy or monitoring; and a child-safety owner who can audit incidents. A classifier score or an LLM self-check is not an escalation service.",
        fill=RISK,
        color="9B1C1C",
    )
    add_body(
        doc,
        "The recommendations below align with child-centred AI principles (development, inclusion, privacy, safety, transparency and accountability) [1]; evidence-based maths guidance on explicit instruction, mathematical language and useful representations [2][3]; and a risk-management lifecycle rather than a one-time model test [10].",
    )

    add_heading(doc, "2. The non-negotiable interaction contract", 1)
    add_bullet(doc, "Truth before fluency. The tutor may say 'I am not sure what you mean yet'; it must not give a polished explanation of an assumed topic.")
    add_bullet(doc, "Dignity before correction. A child can use a nickname, a mixed language, an incomplete sentence or an incorrect term. Repair meaning gently; do not shame grammar, accent, pronunciation, pace or prior knowledge.")
    add_bullet(doc, "Task-specific adaptation, never an IQ judgement. Adapt the explanation to demonstrated understanding in this moment, and let the learner choose the level of detail. Do not expose or infer a permanent ability label.")
    add_bullet(doc, "One primary learning move per turn. Do not answer with five representations at once when the learner says everything feels like a bouncer. Give one small path, then check whether it helped.")
    add_bullet(doc, "Safety supersedes tutoring. A concerning personal disclosure is not a detour to redirect back to maths. Pause the lesson, respond supportively, and follow the appropriate safety route.")
    add_bullet(doc, "Privacy by minimisation. Do not ask for a name, school, address, phone number, password, photo, live location, or story details to explain a lesson or offer basic support.")
    add_bullet(doc, "No pretend capabilities. Never say that a parent, teacher, emergency service, counsellor or human reviewer has been notified unless the product has actually completed that action and can truthfully show its status.")

    add_heading(doc, "3. Required routing model", 1)
    add_body(doc, "Route every utterance through the following order. The steps are deliberately ordered so that a potentially unsafe statement never becomes a normal learning turn or a learner-state update.")
    add_numbered(doc, "Detect high-risk safety, personal-data exposure, harassment/threat, or a request to harm/humiliate somebody. Use high-recall rules plus a trained classifier; the model can add concern but must never remove a deterministic safety flag.")
    add_numbered(doc, "Estimate speech/transcript uncertainty. Preserve alternate interpretations where available. Do not score an answer, change the topic, or update mastery on an uncertain transcription that materially changes meaning.")
    add_numbered(doc, "Resolve the concept only when evidence is sufficient. 'This' and 'that' may use the current on-screen concept only if the reference is explicit and coherent; otherwise ask one short clarification question.")
    add_numbered(doc, "Classify the learning need: purpose/why, definition, procedure, misconception, representation request, practice answer, or frustration/overload.")
    add_numbered(doc, "Choose a learner-controlled level of support, generate one grounded response, and run a separate output verifier before speaking or displaying it.")
    add_numbered(doc, "Ask for a low-effort check ('Did the factor-tree view help, or should I show the division view?') and update only the task-specific learning state supported by the response.")
    add_callout(doc, "Data boundary:", "Safety case data must not be written into mastery, confidence, 'hope', engagement, or future personalisation features. Keep a narrowly-accessible safeguarding record only when policy and law require it; routine analytics should receive a redacted event type, not the child's raw disclosure.", fill=CAUTION, color="7A5A00")

    add_heading(doc, "4. A child asks: 'Why do I have to learn prime factors?'", 1)
    add_heading(doc, "Ground truth that the tutor must get right", 2)
    add_labelled(doc, "Worked fact: ", "3825 = 3 x 3 x 5 x 5 x 17 = 3^2 x 5^2 x 17. A quick independent check is 9 x 25 x 17 = 225 x 17 = 3825.")
    add_labelled(doc, "Precise idea: ", "Every positive whole number greater than 1 can be written as a product of primes in only one way apart from the order of the factors. Prime factorisation reveals that number's building blocks.")
    add_labelled(doc, "Honest usefulness: ", "The student does not need to prime-factorise ordinary numbers every day. The skill is useful inside mathematics for divisibility, HCF/GCD, LCM, simplifying fractions and later algebra/number theory. Primes also matter in some computer-security ideas, but the tutor must not claim that factorising 3825 itself secures anything.")
    add_heading(doc, "Recommended reply for a 15-year-old", 2)
    add_quote(doc, "Fair question. You probably will not need to break every number into primes in daily life. The point is to see a number's recipe - its prime building blocks. For 3825, the recipe is 3 x 3 x 5 x 5 x 17. That helps us compare numbers, find HCF/LCM, simplify fractions, and understand later maths instead of memorising tricks. Want to see the recipe as a factor tree or by repeated division?")
    add_heading(doc, "Why this reply is safe and accurate", 2)
    add_bullet(doc, "It validates the question instead of treating curiosity as resistance.")
    add_bullet(doc, "It does not over-promise a daily real-world use. It names actual mathematical connections and clearly separates a long-term connection to computing from the immediate lesson.")
    add_bullet(doc, "It uses the student's exact number and offers two representations without assuming which one is easier.")
    add_bullet(doc, "It ends with an optional next step, not a forced quiz.")
    add_heading(doc, "Response modes when the same question comes from different learners", 2)
    add_table(
        doc,
        ["Learner signal", "Safe response adjustment", "Do not do"],
        [
            ("'I know this already.'", "Give the short purpose answer, then offer an extension: use the factors to compare 3825 with another number or calculate an HCF/LCM.", "Make them repeat a beginner explanation or infer a high-IQ label."),
            ("'I do not get any word here.'", "Define one term first: 'Prime means a whole number greater than 1 with exactly two positive factors: 1 and itself.' Then return to the purpose.", "Stack definitions for factor, prime, exponent, HCF and LCM in one turn."),
            ("'Everything is a bouncer.'", "Use one concrete move: build 3825 from 25 x 153, then show 25 = 5 x 5 and 153 = 3 x 3 x 17. Pause after each line.", "Switch to a longer, more abstract explanation."),
        ],
        [2200, 4620, 2540],
        font_size=9.3,
    )

    add_heading(doc, "5. Adapt dynamically - but not by IQ", 1)
    add_body(doc, "Yes, the explanation should change dynamically. The design object is a temporary, explainable learning state, not an IQ score. IQ is not needed to decide whether to show an example, define a word, slow down, or offer a challenge; using it risks unfair, sticky and opaque labelling. The learner should see choices such as 'quick version', 'show every step', 'use a picture', 'use a real example', and 'let me try'.")
    add_table(
        doc,
        ["Use as evidence", "May adapt", "Never infer or use"],
        [
            ("A response is correct but the explanation is unsure; a student asks for a definition; an attempt shows a known misconception; learner-selected pace or representation.", "Amount of scaffolding, vocabulary, number of worked steps, representation, practice difficulty, wait time, and whether to ask a check question.", "IQ, disability, trauma history, home situation, race/caste, gender, accent, income, or a single error as a proxy for ability."),
            ("Two or more recent, topic-specific attempts; explicit preference; a student says 'too fast' or 'I want a challenge.'", "Whether to fade a worked example, revisit a prerequisite, or offer extension practice while preserving the same learning objective.", "A permanent 'slow learner' / 'advanced learner' identity or an adaptive profile used for advertising, behavioural monitoring, or non-educational decisions."),
        ],
        [2520, 3320, 3520],
        font_size=9.2,
    )
    add_callout(doc, "Minimum evidence rule:", "Do not lower or raise difficulty after one noisy speech transcript. Require corroborating, topic-specific evidence or a learner choice. Store the reason for the change ('asked for a picture', 'two correct independent attempts'), allow an immediate override, and periodically reset stale inferences.", fill=LIGHT_BLUE, color=INK)
    add_body(doc, "This matches the project's existing direction of adapting from evidence and fading support as mastery is demonstrated. It also aligns with the IES guidance to use explicit, systematic instruction, guided practice, feedback and visual representations rather than a fixed sequence for all students [2][3].")

    add_heading(doc, "6. Generalised handling for 'why', 'what is the use', 'I do not know the word', and overload", 1)
    add_heading(doc, "Purpose questions: a three-part answer, not a slogan", 2)
    add_numbered(doc, "Name the immediate structure: what the idea lets the learner see or do in this chapter.")
    add_numbered(doc, "Name one or two genuine downstream uses in mathematics or science, marked as later applications rather than fake everyday promises.")
    add_numbered(doc, "Offer a choice: a short reason, an example now, or a later application. If the learner says it still feels irrelevant, acknowledge that honestly and return to their current goal.")
    add_quote(doc, "You are allowed to ask that. This idea helps us describe how a change bends rather than only whether it goes up or down. You may not use the word every day, but it is the language used for graphs, motion and later equations. Shall I show one tiny picture first, or explain the word 'curve'?")
    add_heading(doc, "Unknown word: use a micro-definition loop", 2)
    add_numbered(doc, "Give a one-sentence definition in plain language.")
    add_numbered(doc, "Give one example and one non-example when that prevents a common misconception.")
    add_numbered(doc, "Reconnect the word to the current task and ask a one-tap or one-word check.")
    add_quote(doc, "A 'factor' is a whole number that divides another whole number exactly. For example, 5 is a factor of 25 because 25 divided by 5 is 5. Now, which number do you want to factor?")
    add_heading(doc, "'Everything is a bouncer': treat it as overload, not lack of effort", 2)
    add_bullet(doc, "Acknowledge: 'That sounded like too much at once. Let us make it smaller.'")
    add_bullet(doc, "Choose one target: one word, one line, one diagram feature, or one arithmetic move.")
    add_bullet(doc, "Use the learner's preferred modality where known, but present only one primary representation at a time.")
    add_bullet(doc, "After the micro-step, ask 'Did that part make sense: yes, partly, or not yet?' Do not ask 'Do you understand?' with no way to say what failed.")

    add_heading(doc, "7. 'Show it another way' and representation shifts", 1)
    add_body(doc, "A request for another way is meaningful feedback. It usually means the current representation did not connect, not that the child needs more words. The tutor should preserve the same mathematical invariant while changing one representation at a time. Mathematical language and well-chosen concrete or semi-concrete representations are explicitly supported in the IES practice guidance [2].")
    add_table(
        doc,
        ["Representation", "What changes", "What must remain stable"],
        [
            ("Everyday structure", "A fair, non-fake story or grouping situation.", "The same quantities and operation; do not make a story that changes the mathematics."),
            ("Concrete / visual", "Objects, factor tree, number line, graph, table or area model.", "Labels, units, direction and the relationship being shown."),
            ("Symbolic", "Equation, notation, factorisation, graph rule or formal definition.", "Each symbol maps back to a named part of the visual or example."),
            ("Procedural", "A sequence of small reversible steps.", "Why the step is allowed; no unexplained rule swapping."),
            ("Verbal / analogy", "Plain-language rephrase or carefully scoped analogy.", "The analogy's limits. Say where it stops matching the real concept."),
        ],
        [1750, 3800, 3810],
        font_size=9.25,
    )
    add_heading(doc, "Required response pattern", 2)
    add_numbered(doc, "Name the previous view: 'We used a factor tree.'")
    add_numbered(doc, "Offer two genuinely different choices, not two wordier explanations of the same procedure.")
    add_numbered(doc, "Translate one element across views: 'This 5 in the tree is the same 5 in 5 x 5.'")
    add_numbered(doc, "Ask which view helped and record only that preference, not a fixed ability judgement.")
    add_callout(doc, "Representation safety check:", "Never show a visual merely because it matches the broad chapter. It must illustrate the student's exact claim. A wrong diagram can teach a contradiction more confidently than no diagram; abstain and ask or use text-only when relevance is uncertain.", fill=CAUTION, color="7A5A00")

    add_heading(doc, "8. Vague references, half knowledge, wrong terms and wrong pronunciation", 1)
    add_heading(doc, "When the child says 'this' or 'that' without naming the topic", 2)
    add_body(doc, "Do not silently inherit the last concept merely because it is available in session memory. The tutor may use recent on-screen context only when it can state its assumption in the reply. If the reference could point to more than one thing, ask one lightweight question.")
    add_table(
        doc,
        ["Context confidence", "Tutor action", "Example"],
        [
            ("High", "State the assumption and answer narrowly.", "'Do you mean the second step in the factor tree? That step splits 153 into 3 x 51.'"),
            ("Medium", "Offer two short choices based on visible/recent context.", "'Is 'this' the word prime, the tree, or the division step?'"),
            ("Low or conflicting", "Do not pick a concept. Request a small safe fragment.", "'I want to get this right. Please say one word from the question, or paste the sentence. Cover your name or school if it is in a photo.'"),
        ],
        [1700, 3320, 4340],
        font_size=9.25,
    )
    add_heading(doc, "Informal language is useful evidence", 2)
    add_quote(doc, "Calling a parabola a U-shape is a useful start. y = x^2 opens upward like a U, while y = -x^2 opens downward like an upside-down U. Both are parabolas. The equation also tells us exactly how wide it is and where it sits, which the nickname alone cannot tell us.")
    add_body(doc, "The tutor should then invite the child to point to one feature: 'Which way is this graph opening?' This maps informal language to the formal term without rejecting the learner's mental model. It must not imply that every U-shape is a parabola; a parabola has a specific mathematical definition.")
    add_heading(doc, "Pronunciation, grammar and broken sentence formation", 2)
    add_bullet(doc, "Infer meaning from the whole turn and nearby curriculum context, but keep alternate hypotheses when a word changes the maths.")
    add_bullet(doc, "Repair gently: 'I may have heard factors. Did you mean factors or fractions?' A yes/no/tap choice is easier than making the child repeat a long sentence.")
    add_bullet(doc, "Use the correct formal word naturally in the answer. Offer pronunciation only if the child asks, and never make pronunciation a prerequisite for learning the concept.")
    add_bullet(doc, "Do not correct ordinary grammar, accent or code-switching. Correct a mathematical meaning only when needed, and say what changed: 'If you meant denominator rather than numerator, the answer changes because...'")
    add_bullet(doc, "A child can say 'ulta U', 'that bend thing' or a local-language equivalent. Treat it as an invitation to map meanings, not as nonsense.")

    add_heading(doc, "9. Speech-to-text (STT) handling requirements", 1)
    add_body(doc, "The transcript is an estimate of speech, not the child's ground truth. Mathematics is especially sensitive: 'root two' and 'route two', 'factor' and 'fraction', 'x squared' and 'x cube', minus signs, and local-language pronunciation can produce materially different meanings.")
    add_table(
        doc,
        ["Requirement", "Operational rule", "Pass condition"],
        [
            ("Uncertainty preservation", "Carry confidence and, where feasible, N-best hypotheses through intent, concept and grading. Never overwrite the original audio/transcript provenance.", "A low-confidence word that changes an answer produces a clarification, not a score."),
            ("Math-aware normalisation", "Normalise spoken forms only through a constrained maths grammar and the active concept. Keep an auditable parse, e.g., 'three squared' -> 3^2.", "The parser refuses an ambiguous or out-of-grammar rewrite."),
            ("Confirmation before consequence", "Confirm before marking, changing level, recording a misconception, or sending a safety escalation based on uncertain language.", "The child can answer yes/no, tap a displayed alternative, or type."),
            ("Accessible repair", "Offer 'say it again', 'tap the word', 'type it', and 'show the question' paths. Do not require fluent speech.", "Every voice-only failure has a non-voice fallback."),
            ("Privacy", "Treat voice recordings, transcripts and speaker-related data as personal data. Keep the minimum needed and do not use raw recordings for unrelated model improvement without a valid, reviewed basis.", "Retention, access and deletion are documented per locale and age group."),
        ],
        [2000, 4650, 2710],
        font_size=9.1,
    )
    add_quote(doc, "I heard 'three squared'. Did you mean 3^2? Tap Yes or No. If not, you can say the word again or type it.")

    add_heading(doc, "10. Off-topic requests, insults and requests to curse someone", 1)
    add_body(doc, "Do not generate insults, humiliation, threats, harassment or targeted curse messages. The aim is neither to punish the child nor to force an immediate maths redirect. Offer a safer way to express the underlying feeling, and classify bullying, threats or fear of harm separately from ordinary conflict.")
    add_heading(doc, "Recommended response", 2)
    add_quote(doc, "I cannot help write something to hurt or humiliate your friend. If you are angry, you could say: 'I am upset about what happened. Please stop talking to me like that.' If you feel unsafe or someone is bullying you, tell a trusted adult such as a teacher or caregiver.")
    add_bullet(doc, "If the child asks for a joke or a mild comeback, do not guess the social context. Keep the boundary: no targeted abusive content; offer assertive, non-insulting wording.")
    add_bullet(doc, "If there is a threat, coercion, stalking, blackmail, sexual content, violence or fear, route to the safeguarding path. Do not help plan retaliation or investigate the story.")
    add_bullet(doc, "Do not save conflict text as a personality trait, 'aggression score', or engagement signal.")

    add_heading(doc, "11. Personal information and privacy", 1)
    add_body(doc, "A child should never need to disclose personal information to get help with a concept. Data protection is a child-safety requirement, not a settings-page afterthought. UNICEF includes data/privacy and safety among its child-centred AI requirements [1]. UNESCO calls for data privacy protection and age-appropriate human-centred design in educational GenAI [4].")
    add_heading(doc, "Student-facing handling", 2)
    add_quote(doc, "You do not need to share your full name, address, phone number, school, password, code, photo, or live location for me to help with homework. Please keep those private. You can tell me only the maths question.")
    add_body(doc, "If a child has already disclosed something sensitive, do not repeat it back, ask for more identifying details, or make a false promise that it was deleted. Respond to immediate safety first where relevant; otherwise give a simple correction and use the product's truthful privacy/deletion flow.")
    add_heading(doc, "Product requirements", 2)
    add_bullet(doc, "Detect and redact obvious identifiers before ordinary analytics, telemetry, prompts, screenshots and tutor-visible summaries. Preserve only a minimised, access-controlled case reference where a legal/safeguarding process requires review.")
    add_bullet(doc, "Do not use student conversation, voice or safety data for behavioural advertising, targeted advertising, general model training, or cross-context profiling by default.")
    add_bullet(doc, "Make the child-facing notice short and concrete; make the guardian/school notice complete; provide practical access, correction and deletion routes.")
    add_bullet(doc, "Separate learning progress from safety records. A teacher needs a learning summary, not a trauma disclosure; a safety reviewer needs only the minimum material necessary for the case.")
    add_bullet(doc, "Age assurance, parental consent, school authorisation, recording, retention, human review and cross-border transfer must be validated by counsel in every deployment jurisdiction.")
    add_callout(doc, "India deployment anchor:", "India's DPDP Act defines a child as a person under 18 and provides for verifiable parental/guardian consent before processing a child's personal data, with limits on detrimental processing, tracking/behavioural monitoring and targeted advertising [8]. The final deployment must use the current applicable rules and legal review, not this summary as legal advice.", fill=CAUTION, color="7A5A00")
    add_body(doc, "For US child-directed services, COPPA may impose notice, verifiable parental-consent, access, security, retention and deletion duties; it includes audio recordings and persistent identifiers within personal information [9]. These are jurisdictional examples, not a complete compliance list.")

    add_heading(doc, "12. Feelings, trauma and possible immediate danger", 1)
    add_body(doc, "A caring tone matters, but a generic emotional reply is not enough. The tutor must not diagnose, counsel, investigate, promise secrecy, ask the child to recount traumatic details, tell them to confront an alleged abuser, or bounce them back to a lesson before the safety need is assessed. WHO psychological-first-aid guidance emphasises humane, practical help, listening without forcing the person to talk, and protecting people from further harm [5].")
    add_table(
        doc,
        ["Signal", "Tutor's immediate response", "System action"],
        [
            ("Ordinary learning frustration: 'I am bad at maths.'", "Validate and reduce the task: 'This is hard right now, not a verdict on you. We can take one small step or pause.'", "Do not flag as a crisis by default. Keep it out of permanent ability labels."),
            ("Personal distress: sadness, fear, bullying, loss, or 'things are bad at home.'", "Listen briefly; do not press for details. Offer a safe trusted adult not assumed to be a parent. Ask whether they are safe only when the wording indicates risk.", "Route to a trained-human review workflow according to approved policy; show local support resources only from a verified locale configuration."),
            ("Possible self-harm/suicide, violence, abuse, coercion, or immediate danger.", "Pause tutoring. Use a calm, direct safety check and urge immediate human help. Do not continue the academic topic.", "High-priority escalation to a staffed safety team under a reviewed protocol. Record delivery/hand-off status; never claim a notification occurred if it did not."),
            ("Child says the caregiver is the source of harm or may not be safe.", "Do not say 'tell your parent.' Suggest a safe adult who is not the person causing harm: teacher, school counsellor, health worker, trusted relative or another safe local adult.", "Use the child-protection escalation policy and jurisdiction-specific mandatory-reporting advice approved by counsel and safeguarding leadership."),
        ],
        [2200, 4060, 3100],
        font_size=9.0,
    )
    add_heading(doc, "High-risk scripted response requirements", 2)
    add_bullet(doc, "Start with acknowledgement: 'Thank you for telling me. Your safety matters more than the lesson.'")
    add_bullet(doc, "Ask one direct, plain-language question when warranted: 'Are you in immediate danger right now, or thinking about hurting yourself right now?' Do not use euphemisms that hide the safety check.")
    add_bullet(doc, "If yes/unclear: encourage the child, if they can do so safely, to go to a nearby trusted adult or safer shared place and contact emergency help. Keep language short; do not overwhelm them with a long resource list.")
    add_bullet(doc, "If in India, a verified configuration may present the national emergency number 112 for imminent danger and Tele-MANAS (14416 or 1800-89-14416) for mental-health support [6][7]. Outside India, present verified local services; never infer a crisis number from language or IP address alone.")
    add_bullet(doc, "For possible abuse, say they do not need to share more details with the tutor and that they deserve safety. Help them identify a safe adult who is not the alleged source of harm. Do not direct a child to confront anyone.")
    add_bullet(doc, "End only after a designed next step or hand-off. The system may remain available with simple safety-focused prompts, but it must not pretend to be a therapist or substitute for emergency services.")
    add_callout(doc, "Critical distinction:", "Automatic guardian notification is not always safe and must not be the default response to every safety phrase. The escalation protocol needs risk type, child context, jurisdiction, approved contacts, a safe-notification assessment, trained human review and a documented exception process. Generic 'notify the supervising adult now' logic is insufficient.", fill=RISK, color="9B1C1C")

    add_heading(doc, "13. Grounding and verification: how every answer earns the right to be shown", 1)
    add_body(doc, "For a child product, 'the model checked itself' is not verification. Use at least three independent controls for material claims, worked mathematics and safety responses.")
    add_table(
        doc,
        ["Control", "What it checks", "Examples"],
        [
            ("1. Authoritative source or deterministic truth", "Curriculum concept, approved definition, calculation, procedure, age/local resource, and policy template.", "Prime factorisation checked by a deterministic arithmetic verifier; emergency numbers drawn from an approved, versioned locale directory."),
            ("2. Independent response verifier", "Factual consistency, concept match, step validity, misconception match, representation relevance, reading level, unsupported claims, privacy leakage and safety-policy compliance.", "Reject a reply that says every U-shape is a parabola, uses a mismatched diagram, or claims a human was alerted without an action receipt."),
            ("3. Human-reviewed regression suite", "Known misconceptions, STT corruptions, code-switching, vague references, adversarial safety phrasing, privacy leaks and unsafe follow-up behaviour.", "A safeguarding lead and subject expert approve templates and test results; release is blocked if a high-risk case regresses."),
        ],
        [2100, 3570, 3690],
        font_size=9.1,
    )
    add_heading(doc, "Minimum output-verifier checks", 2)
    add_bullet(doc, "Every numerical answer includes a recomputable result; factorisation can be multiplied back to the original number.")
    add_bullet(doc, "Every definition comes from a curated curriculum record or approved source and includes conditions that prevent an overbroad claim.")
    add_bullet(doc, "Every example preserves the concept; an analogy is labelled and its limit is not hidden.")
    add_bullet(doc, "Every context-dependent reply names its assumption or requests clarification.")
    add_bullet(doc, "Every safety reply is selected from reviewed templates keyed to risk type and locale; it cannot be freely improvised by the tutoring model.")
    add_bullet(doc, "Every external resource, helpline and escalation action has an owner, effective date, locale, expiry/review date and verification status.")

    add_heading(doc, "14. Integration notes for the current tutor architecture", 1)
    add_body(doc, "The current project already has useful foundations: a front-door deterministic safety/non-learning gate before cognitive-state updates; an abstention/relevance-floor philosophy; an evidence-driven adaptive practice ladder; and a scripted safety path. These should be retained. The following are required changes before treating the product as child-safe at production level.")
    add_table(
        doc,
        ["Current design area", "Keep", "Required upgrade"],
        [
            ("Front-door routing", "Safety before learning-state changes; model cannot remove a deterministic safety flag.", "Split the broad SAFETY route into personal data, ordinary distress, harassment/threat, abuse/violence, imminent danger, and uncertain-STT safety. Each needs different language and case handling."),
            ("Single scripted SAFETY reply", "Use human-reviewed scripts rather than free model generation.", "Replace the one generic reply with a versioned template library keyed to risk tier, locale, child's safety context and hand-off state. Test every template with safeguarding experts."),
            ("Safety logging and adult notification", "A high-priority incident must not be lost because a tutor reply failed.", "Do not put raw disclosures in general learning logs or learner state. Use restricted, minimised case records; record actual notification delivery; assess whether a named guardian/supervisor is safe before contact."),
            ("Concept inheritance / recent context", "Abstention and relevance floors are the right pattern.", "Do not map ambiguous 'this/that' to INHERIT_CURRENT_CONCEPT without a coreference confidence check. Ask a short clarification when topic evidence is insufficient."),
            ("Adaptive practice", "Evidence-driven fading and learner-state reasoning.", "Expose learner controls, require topic-specific evidence, offer a reset, and prohibit sensitive-profile inputs or permanent IQ-style labels."),
            ("Speech and grading", "STT-aware answer normalisation.", "Carry uncertainty through the entire decision chain and require confirmation before scoring, difficult branching, safety escalation or state updates caused by a disputed transcript."),
        ],
        [1950, 3040, 4370],
        font_size=8.85,
    )
    add_callout(doc, "Important audit finding:", "A test result showing very high recall on a controlled SAFETY dataset is encouraging but is not proof that real children will be protected. Production safety assurance requires open-ended, multilingual, indirect and context-dependent evaluations; monitored failures; trained human operations; privacy controls; and regular independent review.", fill=CAUTION, color="7A5A00")

    add_heading(doc, "15. Test plan and launch gates", 1)
    add_heading(doc, "Required evaluation sets", 2)
    add_bullet(doc, "Purpose/why cases across concepts: includes skeptical, curious, advanced and overwhelmed versions. Verify that uses are true, not invented, and that the tutor does not shame the question.")
    add_bullet(doc, "Definition and representation cases: word unknown, multiple misconceptions, partial mental models, informal labels, diagrams that look similar but encode different concepts, and requests for 'another way'.")
    add_bullet(doc, "Ambiguity/STT cases: broken grammar, local language/code-switching, accents, homophones, math symbols spoken aloud, negative signs, exponents, uncertain words, false starts and interruptions.")
    add_bullet(doc, "Social/off-topic cases: insults, revenge requests, cyberbullying, threats, sexualised content, jokes and ordinary frustration. Confirm the tutor sets a boundary without escalating harmless content into a crisis.")
    add_bullet(doc, "Privacy cases: names, phone numbers, address/location, school, passwords, images and voice. Confirm no unnecessary echoing, logging, training use or displayed retention claim.")
    add_bullet(doc, "Safeguarding cases: indirect and direct disclosures, acute danger, self-harm language, abuse by a caregiver, bullying, grief and false positives. Review the entire multi-turn conversation, not only the first reply.")
    add_heading(doc, "Release gates", 2)
    add_table(
        doc,
        ["Gate", "Evidence required before release"],
        [
            ("Content truth", "Subject-matter review plus deterministic checks for numerical/procedural answers; documented source version for every canonical explanation and local resource."),
            ("Safety operations", "Named safeguarding owner, trained and staffed review/escalation process, service-level expectations, incident playbook, after-hours plan, child-safe notification policy and drill evidence."),
            ("Privacy", "Data map, age/consent review, retention/deletion design, access control, redaction tests, vendor review, no behavioural advertising/profiling decision, and legal sign-off for each target jurisdiction."),
            ("Model and UX", "Independent red-team test, accessibility review, multilingual/STT failure testing, child/educator usability study, override and appeal paths, and a no-blame incident-report mechanism."),
            ("Monitoring", "Versioned templates/resources; dashboards for false negatives/positives and repair loops; a process to pause a harmful response path quickly; periodic re-validation after model, prompt, curriculum or locale changes."),
        ],
        [2240, 7120],
        font_size=9.25,
    )
    add_callout(doc, "Stop-ship condition:", "If the team cannot provide a truthful, staffed hand-off for high-risk disclosures, remove any claim that the tutor monitors safety or alerts adults. The product must still respond supportively and point to verified immediate human help, but it must not simulate a safeguarding service it does not operate.", fill=RISK, color="9B1C1C")

    add_heading(doc, "16. Implementation checklist", 1)
    checks = [
        "Create a risk taxonomy and reviewed template pack for: personal information, learning frustration, emotional distress, bullying/harassment, abuse/violence, self-harm/suicide, imminent danger and ordinary off-topic chat.",
        "Build a locale and resource registry with an owner, validation date, expiry date and an emergency fallback. Do not hard-code numbers in model prompts.",
        "Add an STT uncertainty contract: N-best/confidence, maths parsing, confirmation UI and non-voice fallback before any consequential action.",
        "Add coreference confidence and a clarification UI for 'this/that'; do not use session context as proof of topic identity.",
        "Replace IQ-style personalisation with learner-selected modes and evidence-backed, reversible, topic-specific adaptation.",
        "Create a protected safety case store separate from learning analytics; redact routine logs and restrict safety access by role and need-to-know.",
        "Add a separate response verifier and deterministic maths checks; require three controls for canonical answers and safety templates.",
        "Run child-safety, privacy, accessibility, subject-matter and operations sign-off; repeat whenever a model, prompt, source, policy or locale directory changes.",
    ]
    for item in checks:
        add_bullet(doc, item)

    add_heading(doc, "17. Source notes and grounding", 1)
    add_body(doc, "Sources were checked on 28 July 2026. They ground the design principles and operational cautions; they do not replace local legal counsel, clinical supervision or safeguarding policy. Citation numbers in the document refer to this list.", italic=True)
    references = [
        ("[1] UNICEF (2021). Policy guidance on AI for children, Version 2.0.", "https://www.unicef.org/globalinsight/reports/policy-guidance-ai-children"),
        ("[2] U.S. Department of Education, What Works Clearinghouse (2021). Assisting Students Struggling with Mathematics: Intervention in the Elementary Grades.", "https://ies.ed.gov/ncee/wwc/practiceguide/26"),
        ("[3] U.S. Department of Education, What Works Clearinghouse (2009). Assisting Students Struggling with Mathematics: Response to Intervention for Elementary and Middle Schools.", "https://ies.ed.gov/ncee/wwc/PracticeGuide/2/Published"),
        ("[4] UNESCO (2023, page updated 2026). Guidance for generative AI in education and research.", "https://www.unesco.org/en/articles/guidance-generative-ai-education-and-research"),
        ("[5] World Health Organization, War Trauma Foundation and World Vision International (2011). Psychological first aid: Guide for field workers.", "https://www.who.int/publications-detail-redirect/9789241548205"),
        ("[6] Directorate General of Health Services, Government of India. National Mental Health Programme / Tele-MANAS.", "https://www.dghs.mohfw.gov.in/national-mental-health-programme.php"),
        ("[7] Government of India. Emergency Response Support System (ERSS): 112.", "https://112.gov.in/"),
        ("[8] Government of India (2023). Digital Personal Data Protection Act, 2023, section 9; plus current DPDP Rules materials.", "https://www.meity.gov.in/static/uploads/2024/02/Digital-Personal-Data-Protection-Act-2023.pdf"),
        ("[9] U.S. Federal Trade Commission. Complying with COPPA: Frequently Asked Questions.", "https://www.ftc.gov/business-guidance/resources/complying-coppa-frequently-asked-questions"),
        ("[10] U.S. National Institute of Standards and Technology. AI Risk Management Framework (AI RMF) / Core functions: Govern, Map, Measure, Manage.", "https://www.nist.gov/itl/ai-risk-management-framework"),
    ]
    for text, url in references:
        p = doc.add_paragraph()
        format_paragraph(p, after=4, line=1.08)
        r = p.add_run(text + " ")
        set_run_font(r, size=10)
        add_hyperlink(p, "Open source", url)

    doc.core_properties.title = "AI Tutor: Child-Safe Interaction and Safeguarding Specification"
    doc.core_properties.subject = "Product requirements for safe child tutoring conversations"
    doc.core_properties.author = "AI Tutor Product Team"
    doc.core_properties.comments = "Prepared for internal product and safeguarding review."
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
